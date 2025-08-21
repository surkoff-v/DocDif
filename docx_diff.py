#!/usr/bin/env python3
# pip install mammoth beautifulsoup4 rapidfuzz diff-match-patch python-docx regex
# (опционально, если хотите ещё и Markdown)
# pip install markdownify


import argparse, json, pathlib, regex as re
from dataclasses import dataclass
from typing import List, Tuple, Dict, Optional

import mammoth
from bs4 import BeautifulSoup, NavigableString, Tag
from rapidfuzz import fuzz, process
from diff_match_patch import diff_match_patch
from docx import Document
from docx.shared import RGBColor

WS = re.compile(r'\s+')

def norm(txt: str) -> str:
    return WS.sub(' ', (txt or '').replace('\xa0', ' ')).strip()

@dataclass
class Block:
    idx: int
    section: str
    kind: str   # p|li|th|td|caption|h1|h2|h3...
    text: str

def docx_to_html(path: str) -> str:
    with open(path, 'rb') as f:
        return mammoth.convert_to_html(f).value

def html_to_blocks(html: str) -> List[Block]:
    soup = BeautifulSoup(html, 'html.parser')
    blocks: List[Block] = []
    hstack: List[str] = []

    def section_path() -> str:
        return " > ".join(hstack)

    idx = 0
    def add_block(kind: str, text: str):
        nonlocal idx
        t = norm(text)
        if t:
            blocks.append(Block(idx=idx, section=section_path(), kind=kind, text=t))
            idx += 1

    for el in soup.recursiveChildGenerator():
        if isinstance(el, Tag):
            tag = el.name.lower()
            if tag in ('h1','h2','h3','h4'):
                # обновляем стек секций
                level = int(tag[1])
                while len(hstack) >= level:
                    hstack.pop()
                hstack.append(norm(el.get_text(' ')))
                add_block(tag, el.get_text(' '))
            elif tag in ('p','li','th','td','caption'):
                add_block(tag, el.get_text(' '))
    return blocks

def build_index_text(blocks: List[Block]) -> List[str]:
    return [b.text for b in blocks]

def align_blocks(old: List[Block], new: List[Block], sim_threshold: int = 85
                ) -> Tuple[Dict[int,int], List[int], List[int], Dict[int,int]]:
    """
    Возвращает:
    - matches: old_idx -> new_idx
    - old_only: индексы старых, без пары (delete)
    - new_only: индексы новых, без пары (insert)
    - sim_scores: old_idx -> similarity (0..100) для совпавших
    """
    matches: Dict[int,int] = {}
    sim_scores: Dict[int,int] = {}
    new_available = set(range(len(new)))

    # 1) точные совпадения
    txt2idx_new: Dict[str, List[int]] = {}
    for j, b in enumerate(new):
        txt2idx_new.setdefault(b.text, []).append(j)
    for i, b in enumerate(old):
        js = txt2idx_new.get(b.text)
        if js:
            j = js.pop(0)
            matches[i] = j
            sim_scores[i] = 100
            if not js:
                txt2idx_new.pop(b.text, None)
            if j in new_available:
                new_available.remove(j)

    # 2) «похожие» (RapidFuzz)
    old_unmatched = [i for i in range(len(old)) if i not in matches]
    new_unmatched = list(new_available)
    new_choices = {j: new[j].text for j in new_unmatched}

    for i in old_unmatched:
        if not new_choices: break
        query = old[i].text
        # ищем ближайший кандидат
        best = process.extractOne(
            query, new_choices, scorer=fuzz.token_set_ratio, score_cutoff=sim_threshold
        )
        if best:
            candidate_text, score, j = best
            matches[i] = j
            sim_scores[i] = int(score)
            new_choices.pop(j, None)

    # считаем хвосты
    old_only = [i for i in range(len(old)) if i not in matches]
    new_matched = set(matches.values())
    new_only = [j for j in range(len(new)) if j not in new_matched]
    return matches, old_only, new_only, sim_scores

def tokenize_words(s: str) -> List[str]:
    # Токенизация «по словам/знакам», чтобы дифф был не по символам.
    tokens = re.findall(r'\w+|[^\w\s]', s, flags=re.UNICODE)
    return tokens

def diff_words(a: str, b: str) -> List[Tuple[str,str]]:
    """
    Возвращает список кортежей (op, text), где op: 'eq'|'del'|'ins'
    """
    # Переводим в «строку из токенов» для dmp
    A = tokenize_words(a); B = tokenize_words(b)
    SEP = '\uF000'  # редкий разделитель
    dmp = diff_match_patch()
    diffs = dmp.diff_main(SEP.join(A), SEP.join(B))
    dmp.diff_cleanupSemantic(diffs)

    out: List[Tuple[str,str]] = []
    for op, s in diffs:
        parts = [p for p in s.split(SEP) if p != '']
        if not parts: continue
        text = ' '.join(parts)
        if op == 0: out.append(('eq', text))
        elif op == -1: out.append(('del', text))
        else: out.append(('ins', text))
    return out

def render_html(old_blocks: List[Block], new_blocks: List[Block],
                matches: Dict[int,int], old_only: List[int], new_only: List[int],
                sim_scores: Dict[int,int], out_path: str):
    rows = []
    def mark(parts: List[Tuple[str,str]]):
        left, right = [], []
        for op, t in parts:
            if op == 'eq': left.append(t); right.append(t)
            if op == 'del': left.append(f'<del>{t}</del>')
            if op == 'ins': right.append(f'<ins>{t}</ins>')
        return ' '.join(left), ' '.join(right)

    # заменённые/совпавшие
    for i, j in sorted(matches.items()):
        ob, nb = old_blocks[i], new_blocks[j]
        if ob.text == nb.text:
            # ничего не рисуем? всё равно покажем как "eq"
            parts = [('eq', ob.text)]
        else:
            parts = diff_words(ob.text, nb.text)
        left, right = mark(parts)
        sec = ob.section or nb.section
        sim = sim_scores.get(i, 0)
        rows.append(f"""
        <tr>
          <td><small>old #{i} • {sec} • {ob.kind} • sim:{sim}</small><div>{left}</div></td>
          <td><small>new #{j} • {sec} • {nb.kind} • sim:{sim}</small><div>{right}</div></td>
        </tr>""")

    # удалённые
    for i in old_only:
        ob = old_blocks[i]
        left = f'<del>{ob.text}</del>'
        rows.append(f"""
        <tr>
          <td><small>old #{i} • {ob.section} • {ob.kind}</small><div>{left}</div></td>
          <td><small>—</small><div></div></td>
        </tr>""")

    # добавленные
    for j in new_only:
        nb = new_blocks[j]
        right = f'<ins>{nb.text}</ins>'
        rows.append(f"""
        <tr>
          <td><small>—</small><div></div></td>
          <td><small>new #{j} • {nb.section} • {nb.kind}</small><div>{right}</div></td>
        </tr>""")

    html = f"""<!doctype html><html lang="ru"><head><meta charset="utf-8">
    <title>DOCX diff</title>
    <style>
    body{{font-family:system-ui,Segoe UI,Roboto,Arial;}}
    table{{width:100%;border-collapse:collapse}}
    td{{vertical-align:top;width:50%;border:1px solid #ddd;padding:10px}}
    del{{background:#ffecec;text-decoration:line-through}}
    ins{{background:#eaffea;text-decoration:none}}
    small{{color:#666}}
    </style></head><body>
    <h2>Сравнение DOCX</h2>
    <table>{''.join(rows)}</table>
    </body></html>"""
    pathlib.Path(out_path).write_text(html, encoding='utf-8')

def build_changes_json(old_blocks: List[Block], new_blocks: List[Block],
                       matches: Dict[int,int], old_only: List[int], new_only: List[int],
                       sim_scores: Dict[int,int]) -> List[dict]:
    changes = []
    for i, j in sorted(matches.items()):
        a, b = old_blocks[i], new_blocks[j]
        if a.text != b.text:
            dif = diff_words(a.text, b.text)
            # найдём изменённые числа/даты/проценты
            nums_old = set(re.findall(r'\d+[.,]?\d*%?', a.text))
            nums_new = set(re.findall(r'\d+[.,]?\d*%?', b.text))
            if nums_old != nums_new:
                num_change = {'old': sorted(nums_old), 'new': sorted(nums_new)}
            else:
                num_change = None
            changes.append({
                'type': 'replace',
                'old_idx': i, 'new_idx': j,
                'section': a.section or b.section,
                'similarity': sim_scores.get(i, 0),
                'old_text': a.text, 'new_text': b.text,
                'token_diff': dif,
                'numbers_changed': num_change
            })
    for i in old_only:
        a = old_blocks[i]
        changes.append({'type': 'delete','old_idx': i,'section': a.section,'old_text': a.text})
    for j in new_only:
        b = new_blocks[j]
        changes.append({'type': 'insert','new_idx': j,'section': b.section,'new_text': b.text})
    return changes

def render_docx(changes: List[dict], out_path: str):
    """Генерирует .docx с визуальной подсветкой (не Track Changes)."""
    doc = Document()
    doc.add_heading('DOCX Diff (visual)', level=1)

    def add_run(par, text, op):
        run = par.add_run(text)
        if op == 'del':
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.font.strike = True
        elif op == 'ins':
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
            run.underline = True

    for ch in changes:
        if ch['type'] == 'replace':
            par = doc.add_paragraph()
            par.add_run(f"[{ch.get('section','')}] ").bold = True
            for op, t in ch['token_diff']:
                if op == 'eq': par.add_run(t + ' ')
                elif op in ('del','ins'): add_run(par, t + ' ', op)
        elif ch['type'] == 'delete':
            par = doc.add_paragraph()
            add_run(par, ch['old_text'], 'del')
        elif ch['type'] == 'insert':
            par = doc.add_paragraph()
            add_run(par, ch['new_text'], 'ins')
    doc.save(out_path)

def main():
    ap = argparse.ArgumentParser(description="DOCX diff (Python)")
    ap.add_argument('old_docx')
    ap.add_argument('new_docx')
    ap.add_argument('--out', default='out')
    ap.add_argument('--threshold', type=int, default=85,
                    help='Порог похожести блоков (0..100), по RapidFuzz token_set_ratio')
    args = ap.parse_args()

    out_dir = pathlib.Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    old_html = docx_to_html(args.old_docx)
    new_html = docx_to_html(args.new_docx)
    old_blocks = html_to_blocks(old_html)
    new_blocks = html_to_blocks(new_html)

    matches, old_only, new_only, sim_scores = align_blocks(old_blocks, new_blocks, args.threshold)
    changes = build_changes_json(old_blocks, new_blocks, matches, old_only, new_only, sim_scores)

    # HTML
    render_html(old_blocks, new_blocks, matches, old_only, new_only, sim_scores,
                str(out_dir / 'diff.html'))
    # JSON
    (out_dir / 'changes.json').write_text(json.dumps(changes, ensure_ascii=False, indent=2), encoding='utf-8')
    # DOCX (визуальный)
    render_docx(changes, str(out_dir / 'diff.docx'))

    print(json.dumps({
        "html_report": str(out_dir / "diff.html"),
        "json": str(out_dir / "changes.json"),
        "docx_visual": str(out_dir / "diff.docx"),
        "stats": {
            "old_blocks": len(old_blocks), "new_blocks": len(new_blocks),
            "matched": len(matches), "deleted": len(old_only), "inserted": len(new_only)
        }
    }, ensure_ascii=False))

if __name__ == '__main__':
    main()
